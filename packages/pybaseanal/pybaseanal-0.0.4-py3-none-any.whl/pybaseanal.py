import pandas as pd, numpy as np, seaborn as sns, matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import io
import warnings
warnings.filterwarnings('ignore')


def null_Analysis(data_df, null_Cuttoff = 25, headcount = 5):
    print('Total Data Size {0} and total number of columns are {1}'.format((len(data_df)),(len(data_df.columns))))
    null_df = pd.DataFrame(columns=['Feature', 'Total Non-Null Value', 'No. of Null Value', '% NUll Value', 'Data Type'])
    feature = []
    total_val = []
    null_val = []
    per_null_val = []
    data_type = []
    for i in data_df.columns:
        feature.append(i)
        total_val.append(data_df[i].count())
        null_val.append(data_df[i].isnull().sum())
        per_null_val.append((data_df[i].isnull().sum() / len(data_df))*100)
        data_type.append(data_df[i].dtypes)
    null_df['Feature'] = feature
    null_df['Total  Non-Null Value'] = total_val
    null_df['No. of Null Value'] = null_val
    null_df['% NUll Value'] = per_null_val
    null_df['Data Type'] = data_type
    null_df =  null_df.sort_values(by='% NUll Value', axis=0, ascending=False)
    null_df = null_df.reset_index()
    null_df.drop('index', axis=1, inplace=True)
    high_null_column = []
    low_null_val = []
    for i,j in zip (null_df['Feature'], null_df['% NUll Value']):
        if j > null_Cuttoff:
            high_null_column.append(i)
        elif j < null_Cuttoff and j > 0:
            low_null_val.append(i)
    print('Columns with more than {0}% of null value : {1}'.format(null_Cuttoff, high_null_column) )
    print("")
    print('Columns with less than {0}% of null value : {1}'.format(null_Cuttoff, low_null_val) )
    print("")
    print("*************************")
    print("*** Dataframe Summary ***")
    print("*************************")
    for i in range (0, len(null_df)):
        print("Feature : {0} || Data Type : {1} ". format( null_df['Feature'].iloc[i], null_df['Data Type'].iloc[i]))
        print("No. of Null Value : {0} || Total Non-Null Value : {1} || % NUll Value : {2}".format(null_df['No. of Null Value'].iloc[i],null_df['Total  Non-Null Value'].iloc[i],null_df['% NUll Value'].iloc[i]))
        print('===========')


def plot_categorical_frequency(df, output_path='cat_graph.docx'):
    categorical_columns = df.select_dtypes(include='object').columns
    document = Document()

    for column in categorical_columns:
        category_counts = df[column].value_counts()

        if len(category_counts) == 1:
            print(f"{column} has only one category: {category_counts.index[0]} (Total count: {category_counts.iloc[0]})")
            continue

        if len(category_counts) > 10:
            top_categories = category_counts.head(10)
            other_categories_count = category_counts[10:].sum()
            category_counts = top_categories.append(pd.Series({'Others': other_categories_count}))

        plt.figure(figsize=(10, 8))  # Adjusted figure size
        plt.subplots_adjust(bottom=0.2)  # Set bottom margin

        category_counts.plot(kind='bar', color='skyblue')
        plt.title(f'Frequency Distribution of {column}')
        plt.xlabel(column)
        plt.ylabel('Frequency')

        for i, count in enumerate(category_counts):
            plt.text(i, count + 0.1, str(count), ha='center')    

        buffer = io.BytesIO()
        plt.savefig(buffer, format='jpg')
        plt.close()

        document.add_picture(buffer, width=Inches(5), height=Inches(4))  # Adjusted image size

    document.save(output_path)


def convert_columns_to_preferred_types(df):
    for column in df.columns:
        original_dtype = df[column].dtype
        if pd.api.types.is_string_dtype(df[column]):
            try:
                df[column] = pd.to_datetime(df[column])
                print(f"Converted {column} to datetime")
            except ValueError:
                pass
        if pd.api.types.is_object_dtype(df[column]):
            try:
                df[column] = pd.to_numeric(df[column])
                print(f"Converted {column} to numeric")
            except ValueError:
                pass  
        new_dtype = df[column].dtype
        if original_dtype != new_dtype:
            print(f"{column} data type changed from {original_dtype} to {new_dtype}")
    return df



def generate_visualizations_and_report(df, output_path='output.docx'):

    document = Document()

    numerical_columns = df.select_dtypes(include=['number']).columns
    n = 1
    for column in numerical_columns:
        if len(df[column].unique()) == 1:
            print(f"{column} has only one category: {df[column].unique()[0]} (Total count: {len(df)})")
            continue

        if len(df[column].unique()) > 1:
            plt.figure(figsize=(10, 6))
            sns.lineplot(x=df.index, y=df[column])
            plt.title(f'Line Chart for {column}')
            plt.xlabel('Index')
            plt.ylabel(column)
            buffer = io.BytesIO()
            plt.savefig(buffer, format='jpg')
            plt.close()

            document.add_picture(buffer, width=Inches(5))
            plt.show()


        if len(df[column].unique()) <= 10:
            plt.figure(figsize=(10, 6))
            sns.countplot(x=df[column])
            plt.title(f'Bar Chart for {column}')
            plt.xlabel(column)
            plt.ylabel('Count')
            buffer = io.BytesIO()
            plt.savefig(buffer, format='jpg')
            plt.close()
            document.add_picture(buffer, width=Inches(5))
            plt.show()

        if len(df[column].unique()) > 1:
            plt.figure(figsize=(10, 6))
            sns.histplot(df[column], kde=True)
            plt.title(f'Histogram for {column}')
            plt.xlabel(column)
            plt.ylabel('Frequency')
            buffer = io.BytesIO()
            plt.savefig(buffer, format='jpg')
            plt.close()
            document.add_picture(buffer, width=Inches(5))
            plt.show()

        plt.figure(figsize=(10, 6))
        sns.boxplot(x=df[column])
        plt.title(f'Box Plot for {column}')
        plt.xlabel(column)
        plt.ylabel('Values')
        buffer = io.BytesIO()
        plt.savefig(buffer, format='jpg')
        plt.close()
        document.add_picture(buffer, width=Inches(5))
        plt.show()

        Q1 = df[column].quantile(0.25)
        Q3 = df[column].quantile(0.75)
        IQR = Q3 - Q1
        outliers_count = len(df[(df[column] < (Q1 - 1.5 * IQR)) | (df[column] > (Q3 + 1.5 * IQR))])

        stats_table = document.add_table(rows=1, cols=5)
        stats_table.autofit = True
        stats_table.style = 'Table Grid'
        stats_table.cell(0, 0).text = 'Column Name'
        stats_table.cell(0, 1).text = 'Mean'
        stats_table.cell(0, 2).text = 'Median'
        stats_table.cell(0, 3).text = 'Std Dev'
        stats_table.cell(0, 4).text = 'IQR'

        stats_table.add_row()
        stats_table.cell(1, 0).text = column
        stats_table.cell(1, 1).text = f'{df[column].mean():.2f}'
        stats_table.cell(1, 2).text = f'{df[column].median():.2f}'
        stats_table.cell(1, 3).text = f'{df[column].std():.2f}'
        stats_table.cell(1, 4).text = f'{IQR:.2f}'
        print(f"Completed {n}/{len(df.columns)}",  end="\r")
        n += 1
    document.save(output_path)